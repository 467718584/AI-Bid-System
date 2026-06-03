"""
Word文档模板样式导出器
"""
import io
import logging
from typing import Optional, List, Dict, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from bid_templates import BidTemplate, get_template, list_templates, TemplateType

logger = logging.getLogger(__name__)


class StyledDocumentExporter:
    """带样式的Word文档导出器"""
    
    def __init__(self, template: Optional[BidTemplate] = None):
        self.template = template or get_template(TemplateType.STANDARD)
    
    def create_document(self) -> Document:
        """创建带样式的文档"""
        doc = Document()
        
        # 设置页面
        section = doc.sections[0]
        section.page_width = self.template.page_width
        section.page_height = self.template.page_height
        section.top_margin = self.template.margin_top
        section.bottom_margin = self.template.margin_bottom
        section.left_margin = self.template.margin_left
        section.right_margin = self.template.margin_right
        
        # 设置默认样式
        style = doc.styles["Normal"]
        style.font.name = self.template.body_font
        style.font.size = Pt(self.template.body_size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), self.template.body_font)
        
        return doc
    
    def add_title(self, doc: Document, text: str, level: int = 0):
        """添加标题"""
        if level == 0:
            # 文档主标题
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = self.template.title_font
            run.font.size = Pt(self.template.title_size)
            run.font.color.rgb = self.template.title_color
            run.bold = self.template.title_bold
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.template.title_font)
            para.alignment = self.template.title_alignment
            para.space_after = Pt(24)
        elif level == 1:
            para = doc.add_heading(text, level=1)
            self._apply_h1_style(para)
        elif level == 2:
            para = doc.add_heading(text, level=2)
            self._apply_h2_style(para)
        elif level == 3:
            para = doc.add_heading(text, level=3)
            self._apply_h3_style(para)
        else:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(self.template.body_size)
            run.bold = True
        
        return para
    
    def _apply_h1_style(self, para):
        """应用一级标题样式"""
        para.alignment = self.template.h1_alignment
        para.space_before = self.template.h1_space_before
        para.space_after = self.template.h1_space_after
        for run in para.runs:
            run.font.name = self.template.h1_font
            run.font.size = Pt(self.template.h1_size)
            run.font.color.rgb = self.template.h1_color
            run.bold = self.template.h1_bold
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.template.h1_font)
    
    def _apply_h2_style(self, para):
        """应用二级标题样式"""
        para.alignment = self.template.h2_alignment
        para.space_before = self.template.h2_space_before
        para.space_after = self.template.h2_space_after
        for run in para.runs:
            run.font.name = self.template.h2_font
            run.font.size = Pt(self.template.h2_size)
            run.font.color.rgb = self.template.h2_color
            run.bold = self.template.h2_bold
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.template.h2_font)
    
    def _apply_h3_style(self, para):
        """应用三级标题样式"""
        for run in para.runs:
            run.font.name = self.template.h3_font
            run.font.size = Pt(self.template.h3_size)
            run.font.color.rgb = self.template.h3_color
            run.bold = self.template.h3_bold
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.template.h3_font)
    
    def add_paragraph(self, doc: Document, text: str, bold: bool = False,
                      font_size: Optional[int] = None, alignment: str = "LEFT"):
        """添加正文段落"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = self.template.body_font
        run.font.size = Pt(font_size or self.template.body_size)
        run.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.template.body_font)
        
        # 设置对齐
        if alignment == "CENTER":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "RIGHT":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == "JUSTIFY":
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = self.template.body_alignment
        
        # 设置行距
        para.paragraph_format.line_spacing = self.template.body_line_spacing
        para.paragraph_format.first_line_indent = self.template.body_first_line_indent
        
        return para
    
    def add_styled_table(self, doc: Document, data: List[List[Any]],
                         headers: Optional[List[str]] = None,
                         title: Optional[str] = None):
        """添加带样式的表格"""
        if not data:
            return
        
        # 表格标题
        if title:
            caption_para = doc.add_paragraph()
            run = caption_para.add_run(title)
            run.font.size = Pt(11)
            run.bold = True
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        rows_count = len(data) + (1 if headers else 0)
        cols_count = len(data[0]) if data else 0
        
        table = doc.add_table(rows=rows_count, cols=cols_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        if headers:
            header_row = table.rows[0]
            for i, header_text in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header_text
                
                # 设置表头样式
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.runs[0] if para.runs else para.add_run(header_text)
                run.font.name = self.template.table_header_font
                run.font.size = Pt(self.template.table_header_size)
                run.bold = True
                run.font.color.rgb = getattr(self.template, 'table_header_font_color', RGBColor(255, 255, 255))
                run._element.rPr.rFonts.set(qn("w:eastAsia"), self.template.table_header_font)
                
                # 设置背景色
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), self.template.table_header_bg._str if hasattr(self.template.table_header_bg, '_str') else self.template.table_header_bg.__str__())
                cell._tc.get_or_add_tcPr().append(shading)
        
        # 数据行
        start_idx = 1 if headers else 0
        for i, row_data in enumerate(data):
            row = table.rows[i + start_idx]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_text)
                
                para = cell.paragraphs[0]
                if para.runs:
                    run = para.runs[0]
                    run.font.name = self.template.table_body_font
                    run.font.size = Pt(self.template.table_body_size)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), self.template.table_body_font)
                
                # 斑马纹
                if i % 2 == 1:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'F2F2F2')
                    cell._tc.get_or_add_tcPr().append(shading)
        
        return table
    
    def add_html_content(self, doc: Document, html_content: str):
        """从HTML内容添加到文档"""
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for element in soup.children:
            if element.name is None:
                continue
            
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1]) if element.name[1].isdigit() else 1
                text = element.get_text().strip()
                if text:
                    self.add_title(doc, text, level=level)
            
            elif element.name == 'p':
                text = element.get_text().strip()
                if text:
                    strong = element.find('strong')
                    self.add_paragraph(doc, text, bold=strong is not None)
            
            elif element.name == 'div':
                self.add_html_content(doc, str(element))
            
            elif element.name == 'strong':
                text = element.get_text().strip()
                if text:
                    self.add_paragraph(doc, text, bold=True)
            
            elif element.name == 'table':
                self._add_html_table(doc, element)
            
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if text:
                        para = doc.add_paragraph(style='List Bullet')
                        run = para.add_run(text)
                        run.font.name = self.template.body_font
                        run.font.size = Pt(self.template.body_size)
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.template.body_font)
            
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if text:
                        para = doc.add_paragraph(style='List Number')
                        run = para.add_run(text)
                        run.font.name = self.template.body_font
                        run.font.size = Pt(self.template.body_size)
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.template.body_font)
    
    def _add_html_table(self, doc: Document, table_element):
        """从HTML表格添加"""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # 获取表头
        header_row = rows[0] if rows else None
        headers = []
        if header_row:
            for th in header_row.find_all(['th', 'td']):
                headers.append(th.get_text().strip())
        
        # 获取数据
        data_rows = []
        for tr in rows[1:]:
            row_data = []
            for td in tr.find_all('td'):
                row_data.append(td.get_text().strip())
            if row_data:
                data_rows.append(row_data)
        
        if headers or data_rows:
            self.add_styled_table(doc, data_rows, headers if headers else None)
    
    def export_to_bytes(self, doc: Document) -> bytes:
        """导出文档到字节流"""
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
