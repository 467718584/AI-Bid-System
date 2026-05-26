"""Word表格导出模块 - 增强版表格样式支持"""
import io
from typing import Dict, Any, List, Optional, Union

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class TableExporter:
    """Word表格导出器 - 支持合并单元格、背景色、表头样式、边框样式"""

    def __init__(self):
        self.default_font = "宋体"
        self.default_size = 10.5

        # 预定义样式
        self.styles = {
            "header_blue": {
                "bg_color": "4472C4",
                "font_color": "FFFFFF",
                "bold": True,
                "align": "CENTER"
            },
            "header_green": {
                "bg_color": "70AD47",
                "font_color": "FFFFFF",
                "bold": True,
                "align": "CENTER"
            },
            "header_gray": {
                "bg_color": "595959",
                "font_color": "FFFFFF",
                "bold": True,
                "align": "CENTER"
            },
            "data_row": {
                "bg_color": None,
                "font_color": "000000",
                "bold": False,
                "align": "LEFT"
            },
            "alt_row": {
                "bg_color": "D9E2F3",
                "font_color": "000000",
                "bold": False,
                "align": "LEFT"
            },
            "total_row": {
                "bg_color": "B4C6E7",
                "font_color": "000000",
                "bold": True,
                "align": "CENTER"
            },
            "highlight": {
                "bg_color": "FFF2CC",
                "font_color": "000000",
                "bold": False,
                "align": "LEFT"
            }
        }

    def _set_cell_text(self, cell, text: str, font_size: float = 10.5,
                       bold: bool = False, align: str = "LEFT",
                       font_name: str = None, font_color: str = None):
        """设置单元格文本样式"""
        cell.text = text
        para = cell.paragraphs[0]

        # 设置对齐方式
        if align == "CENTER":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "RIGHT":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 设置字体
        run = para.runs[0] if para.runs else para.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = font_name or self.default_font

        if font_color:
            run.font.color.rgb = RGBColor.from_string(font_color)

    def _set_cell_bg_color(self, cell, color: str):
        """设置单元格背景色"""
        if not color:
            return

        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_cell_border(self, cell, border_info: Dict[str, Any]):
        """设置单元格边框

        Args:
            cell: 单元格对象
            border_info: 边框信息，包含:
                - top, bottom, left, right: 边框设置
                每个边框设置包含: val (single/double/none), sz (宽度), color (颜色)
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        for border_name in ['top', 'left', 'bottom', 'right']:
            if border_name in border_info:
                border_data = border_info[border_name]
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), border_data.get('val', 'single'))
                border.set(qn('w:sz'), str(border_data.get('sz', 4)))
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), border_data.get('color', '000000'))

                # 移除已存在的边框元素
                existing = tcPr.find(qn(f'w:{border_name}'))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.append(border)

    def _set_cell_vertical_align(self, cell, align: str = "CENTER"):
        """设置单元格垂直对齐"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), align.lower())  # top, center, bottom
        tcPr.append(vAlign)

    def merge_cells_horizontal(self, table, row_idx: int, start_col: int, end_col: int):
        """水平合并单元格"""
        row = table.rows[row_idx]
        cell = row.cells[start_col]
        for i in range(start_col + 1, end_col + 1):
            cell = cell.merge(row.cells[i])

    def merge_cells_vertical(self, table, start_row: int, end_row: int, col_idx: int):
        """垂直合并单元格"""
        cell = table.rows[start_row].cells[col_idx]
        for i in range(start_row + 1, end_row + 1):
            cell = cell.merge(table.rows[i].cells[col_idx])

    def apply_row_border(self, table, row_idx: int, border_info: Dict[str, Any]):
        """为整行应用边框"""
        row = table.rows[row_idx]
        for cell in row.cells:
            self._set_cell_border(cell, border_info)

    def apply_column_width(self, table, col_idx: int, width: float):
        """设置列宽（英寸）"""
        for row in table.rows:
            cell = row.cells[col_idx]
            cell.width = Inches(width)

    def set_table_border(self, table, border_info: Dict[str, Any]):
        """设置整个表格的边框

        Args:
            border_info: 边框信息，包含:
                - top, bottom, left, right, insideH, insideV
        """
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            if border_name in border_info:
                border_data = border_info[border_name]
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), border_data.get('val', 'single'))
                border.set(qn('w:sz'), str(border_data.get('sz', 4)))
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), border_data.get('color', '000000'))

                # 移除已存在的边框
                existing = tblBorders.find(qn(f'w:{border_name}'))
                if existing is not None:
                    tblBorders.remove(existing)
                tblBorders.append(border)

    def create_table(self, doc: Document, headers: List[str],
                     data: List[List[str]],
                     style_config: Dict[str, Any] = None,
                     column_widths: List[float] = None,
                     merge_info: List[Dict] = None) -> Document:
        """创建带完整样式的表格

        Args:
            doc: Word文档对象
            headers: 表头列表
            data: 数据行列表
            style_config: 样式配置
            column_widths: 列宽列表（英寸）
            merge_info: 合并信息列表 [{"type": "horizontal", "row": 0, "start": 0, "end": 2}]

        Returns:
            Word文档对象
        """
        if style_config is None:
            style_config = {}

        # 创建表格
        rows = len(data) + 1
        cols = len(headers)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        # 设置表格对齐
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 应用外边框样式
        default_border = {
            "top": {"val": "single", "sz": 8, "color": "4472C4"},
            "bottom": {"val": "single", "sz": 8, "color": "4472C4"},
            "left": {"val": "single", "sz": 8, "color": "4472C4"},
            "right": {"val": "single", "sz": 8, "color": "4472C4"},
            "insideH": {"val": "single", "sz": 4, "color": "4472C4"},
            "insideV": {"val": "single", "sz": 4, "color": "4472C4"}
        }
        self.set_table_border(table, default_border)

        # 填充表头
        header_style = style_config.get("header", self.styles["header_blue"])

        for i, header_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            self._set_cell_bg_color(cell, header_style["bg_color"])
            self._set_cell_text(
                cell, header_text,
                bold=True,
                font_size=style_config.get("header_font_size", 10.5),
                align=header_style.get("align", "CENTER"),
                font_color=header_style.get("font_color", "FFFFFF")
            )
            self._set_cell_vertical_align(cell, "CENTER")

        # 填充数据行
        data_style = style_config.get("data_row", self.styles["data_row"])
        alt_style = style_config.get("alt_row", self.styles["alt_row"])
        total_style = style_config.get("total_row", self.styles["total_row"])

        for row_idx, row_data in enumerate(data):
            row = table.rows[row_idx + 1]
            is_total = row_data[0] == "" and any("合计" in str(c) for c in row_data)
            is_alt = row_idx % 2 == 1

            # 确定行样式
            if is_total:
                row_style = total_style
            elif is_alt:
                row_style = alt_style
            else:
                row_style = data_style

            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                self._set_cell_bg_color(cell, row_style.get("bg_color"))
                self._set_cell_text(
                    cell, str(cell_text),
                    bold=row_style.get("bold", False),
                    font_size=style_config.get("data_font_size", 10),
                    align=row_style.get("align", "LEFT") if col_idx > 0 else "CENTER",
                    font_color=row_style.get("font_color", "000000")
                )
                self._set_cell_vertical_align(cell, "CENTER")

        # 应用列宽
        if column_widths:
            for col_idx, width in enumerate(column_widths):
                if col_idx < cols:
                    self.apply_column_width(table, col_idx, width)

        # 处理单元格合并
        if merge_info:
            for merge_item in merge_info:
                if merge_item["type"] == "horizontal":
                    self.merge_cells_horizontal(
                        table,
                        merge_item["row"],
                        merge_item["start"],
                        merge_item["end"]
                    )
                elif merge_item["type"] == "vertical":
                    self.merge_cells_vertical(
                        table,
                        merge_item["start_row"],
                        merge_item["end_row"],
                        merge_item["col"]
                    )

        return table

    def create_qualification_table(self, doc: Document, requirements: List[Dict[str, Any]]) -> Document:
        """创建资质要求对照表"""
        doc.add_heading("二、资质要求对照表", level=2)

        headers = ["资质名称", "要求等级", "必需/可选", "有效期限", "我方状态", "备注"]
        data = []

        for req in requirements:
            status = req.get("status", "")
            if "已满足" in status:
                status_display = f"✓ {status}"
            elif "不满足" in status:
                status_display = f"✗ {status}"
            elif "部分" in status:
                status_display = f"△ {status}"
            else:
                status_display = status

            data.append([
                req.get("name", ""),
                req.get("level", ""),
                req.get("required", "必需"),
                req.get("deadline", ""),
                status_display,
                req.get("remark", "")
            ])

        # 添加统计行
        total_req = len(requirements)
        passed = sum(1 for r in requirements if "已满足" in r.get("status", ""))
        data.append(["", f"合计：{total_req}项（已满足{passed}项）", "", "", "", ""])

        column_widths = [1.8, 1.2, 0.8, 0.9, 0.9, 1.2]
        merge_info = [{"type": "horizontal", "row": len(data), "start": 1, "end": 5}]

        self.create_table(doc, headers, data,
                         style_config={"header": self.styles["header_blue"]},
                         column_widths=column_widths,
                         merge_info=merge_info)

        return doc

    def create_personnel_table(self, doc: Document, positions: List[Dict[str, Any]]) -> Document:
        """创建项目人员配置表"""
        doc.add_heading("三、项目人员配置表", level=2)

        headers = ["序号", "岗位名称", "人数", "专业要求", "资质要求", "主要职责"]
        data = []

        for idx, pos in enumerate(positions, 1):
            data.append([
                str(idx),
                pos.get("name", ""),
                str(pos.get("count", 1)),
                pos.get("major", ""),
                pos.get("qualification", ""),
                pos.get("responsibility", "")
            ])

        # 添加合计行
        total_count = sum(int(p.get("count", 0)) for p in positions)
        data.append(["", f"合计", str(total_count), "", "", ""])

        column_widths = [0.5, 1.0, 0.5, 1.2, 1.2, 2.4]
        merge_info = [{"type": "horizontal", "row": len(data), "start": 1, "end": 5}]

        self.create_table(doc, headers, data,
                         style_config={"header": self.styles["header_green"]},
                         column_widths=column_widths,
                         merge_info=merge_info)

        return doc

    def create_schedule_table(self, doc: Document, phases: List[Dict[str, Any]],
                             project_name: str = "") -> Document:
        """创建施工进度计划表"""
        title = f"四、施工进度计划表"
        if project_name:
            title = f"四、{project_name} - 施工进度计划表"
        doc.add_heading(title, level=2)

        headers = ["序号", "工作阶段", "开始时间", "结束时间", "持续天数", "关键节点", "备注"]
        data = []

        for idx, phase in enumerate(phases, 1):
            data.append([
                str(idx),
                phase.get("name", ""),
                phase.get("start_date", ""),
                phase.get("end_date", ""),
                str(phase.get("duration", "")),
                phase.get("milestone", ""),
                phase.get("remark", "")
            ])

        # 计算总工期
        if phases:
            total_days = sum(phase.get("duration", 0) for phase in phases)
            data.append(["", f"总工期：{total_days}天", "", "", "", "", ""])
            merge_info = [{"type": "horizontal", "row": len(data), "start": 1, "end": 6}]
        else:
            merge_info = None

        column_widths = [0.4, 1.2, 0.9, 0.9, 0.7, 1.2, 0.5]

        self.create_table(doc, headers, data,
                         style_config={"header": self.styles["header_gray"]},
                         column_widths=column_widths,
                         merge_info=merge_info)

        return doc

    def create_equipment_table(self, doc: Document, equipment_list: List[Dict[str, Any]]) -> Document:
        """创建设备清单表"""
        doc.add_heading("五、设备清单表", level=2)

        headers = ["序号", "设备名称", "规格型号", "数量", "来源", "现状", "部署位置"]
        data = []

        for idx, eq in enumerate(equipment_list, 1):
            data.append([
                str(idx),
                eq.get("name", ""),
                eq.get("model", ""),
                str(eq.get("count", 1)),
                eq.get("owner_or_lease", "自有"),
                eq.get("condition", "完好"),
                eq.get("deployment_location", "")
            ])

        # 添加合计行
        total_count = sum(int(eq.get("count", 0)) for eq in equipment_list)
        data.append(["", "合计", "", str(total_count), "", "", ""])

        column_widths = [0.4, 1.2, 1.2, 0.5, 0.6, 0.6, 1.3]
        merge_info = [{"type": "horizontal", "row": len(data), "start": 1, "end": 6}]

        self.create_table(doc, headers, data,
                         style_config={"header": self.styles["header_blue"]},
                         column_widths=column_widths,
                         merge_info=merge_info)

        return doc

    def create_safety_table(self, doc: Document, measures: List[Dict[str, Any]]) -> Document:
        """创建安全措施检查表"""
        doc.add_heading("六、安全措施检查表", level=2)

        headers = ["序号", "类别", "安全措施", "检查标准", "责任人", "检查频率", "备注"]
        data = []

        for idx, m in enumerate(measures, 1):
            data.append([
                str(idx),
                m.get("category", ""),
                m.get("measure", ""),
                m.get("standard", ""),
                m.get("person_responsible", ""),
                m.get("check_frequency", ""),
                m.get("remark", "")
            ])

        column_widths = [0.4, 0.8, 1.4, 1.4, 0.8, 0.7, 0.5]

        self.create_table(doc, headers, data,
                         style_config={"header": self.styles["header_green"]},
                         column_widths=column_widths)

        return doc

    def export_tables_to_word(self, table_data: Dict[str, Any],
                              doc_title: str = "技术标表格") -> bytes:
        """导出所有表格到Word文档

        Args:
            table_data: 包含所有表格数据的字典
            doc_title: 文档标题

        Returns:
            Word文档二进制数据
        """
        doc = Document()

        # 设置默认样式
        style = doc.styles["Normal"]
        style.font.name = self.default_font
        style.font.size = Pt(self.default_size)

        # 添加文档标题
        doc.add_heading(doc_title, level=1)

        # 添加资质要求表
        if "qualifications" in table_data:
            self.create_qualification_table(doc, table_data["qualifications"])

        # 添加人员配置表
        if "personnel" in table_data:
            self.create_personnel_table(doc, table_data["personnel"])

        # 添加进度计划表
        if "schedule" in table_data:
            schedule_data = table_data["schedule"]
            self.create_schedule_table(
                doc,
                schedule_data.get("phases", []),
                schedule_data.get("project_name", "")
            )

        # 添加设备清单表
        if "equipment" in table_data:
            self.create_equipment_table(doc, table_data["equipment"])

        # 添加安全措施表
        if "safety" in table_data:
            self.create_safety_table(doc, table_data["safety"])

        # 保存到字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def add_gantt_chart(self, doc: Document, gantt_image: bytes,
                        title: str = "施工进度计划（甘特图）") -> Document:
        """向文档添加甘特图

        Args:
            doc: Word文档对象
            gantt_image: 甘特图二进制数据
            title: 图表标题

        Returns:
            Word文档对象
        """
        doc.add_heading(title, level=2)

        # 添加甘特图图片
        buffer = io.BytesIO(gantt_image)
        doc.add_picture(buffer, width=Inches(6.5))

        # 添加图注
        para = doc.add_paragraph()
        run = para.add_run("注：图中红色虚线表示当前日期")
        run.font.size = Pt(9)
        run.italic = True

        return doc