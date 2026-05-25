"""文档解析模块 - 支持PDF、Word、政府采购格式"""
import io
import logging
import re
from typing import Dict, Any, Optional, Tuple
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """解析后的文档对象"""
    file_name: str
    file_type: str
    content: str
    basic_info: Dict[str, Any]
    scoring_method: Dict[str, Any]
    compliance_items: list
    disqualification_items: list
    raw_text: str


class DocumentParser:
    """文档解析器"""

    def __init__(self):
        self.supported_formats = ["pdf", "docx", "doc", "zf"]

    async def parse(self, file_content: bytes, file_name: str, file_type: str) -> ParsedDocument:
        """解析文档"""
        if file_type.lower() == "pdf":
            return await self._parse_pdf(file_content, file_name)
        elif file_type.lower() in ["docx", "doc"]:
            return await self._parse_word(file_content, file_name)
        elif file_type.lower() == "zf":
            return await self._parse_zf(file_content, file_name)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    async def _parse_pdf(self, content: bytes, file_name: str) -> ParsedDocument:
        """解析PDF文件"""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.pdf.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            full_text = "\n".join(text_parts)

            return ParsedDocument(
                file_name=file_name,
                file_type="pdf",
                content=full_text,
                basic_info=self._extract_basic_info(full_text),
                scoring_method=self._extract_scoring_method(full_text),
                compliance_items=self._extract_compliance_items(full_text),
                disqualification_items=self._extract_disqualification_items(full_text),
                raw_text=full_text
            )
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise

    async def _parse_word(self, content: bytes, file_name: str) -> ParsedDocument:
        """解析Word文件"""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)

            # 提取表格
            tables = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                tables.append("\n".join(table_text))

            full_text_with_tables = full_text + "\n\n" + "\n\n".join(tables)

            return ParsedDocument(
                file_name=file_name,
                file_type="word",
                content=full_text_with_tables,
                basic_info=self._extract_basic_info(full_text),
                scoring_method=self._extract_scoring_method(full_text),
                compliance_items=self._extract_compliance_items(full_text),
                disqualification_items=self._extract_disqualification_items(full_text),
                raw_text=full_text
            )
        except Exception as e:
            logger.error(f"Word parsing error: {e}")
            raise

    async def _parse_zf(self, content: bytes, file_name: str) -> ParsedDocument:
        """解析政府采购格式（zf实际是特殊编码的文本文件）"""
        try:
            text = ""
            # 尝试多种编码
            for encoding in ["utf-8", "gbk", "gb2312"]:
                try:
                    text = content.decode(encoding)
                    break
                except Exception:
                    continue

            return ParsedDocument(
                file_name=file_name,
                file_type="zf",
                content=text,
                basic_info=self._extract_basic_info(text),
                scoring_method=self._extract_scoring_method(text),
                compliance_items=self._extract_compliance_items(text),
                disqualification_items=self._extract_disqualification_items(text),
                raw_text=text
            )
        except Exception as e:
            logger.error(f"ZF parsing error: {e}")
            raise

    def _extract_basic_info(self, text: str) -> Dict[str, Any]:
        """提取基本信息"""
        info = {}

        # 项目名称（通常在标题或开头）
        title_match = re.search(r"(.+?招标公告|.+?招标文件|项目名称[：:](.+?)\n)", text[:500])
        if title_match:
            info["project_name"] = title_match.group(1).strip()

        # 招标代理
        agency_match = re.search(r"招标代理[：:]\s*([^\n]+)", text)
        if agency_match:
            info["agency_name"] = agency_match.group(1).strip()

        # 联系人
        contact_match = re.search(r"联系人[：:]\s*([^\n]+)", text)
        if contact_match:
            info["contact_person"] = contact_match.group(1).strip()

        # 联系电话
        phone_match = re.search(r"联系电话[：:]\s*([^\n]+)", text)
        if phone_match:
            info["contact_phone"] = phone_match.group(1).strip()

        # 投标截止时间
        deadline_match = re.search(r"投标截止|截止时间[：:]\s*([^\n]+)", text)
        if deadline_match:
            info["bid_deadline"] = deadline_match.group(1).strip()

        return info

    def _extract_scoring_method(self, text: str) -> Dict[str, Any]:
        """提取评标办法"""
        method = {
            "disqualification_items": [],
            "preliminary_review": [],
            "commercial_review": [],
            "technical_review": []
        }

        # 废标条款
        disqual_pattern = r"废标|作无效投标|不符合招标文件"
        for match in re.finditer(disqual_pattern, text):
            start = max(0, match.start() - 100)
            end = min(len(text), match.end() + 100)
            context = text[start:end].replace("\n", " ").strip()
            method["disqualification_items"].append(context)

        # 初步评审
        prelim_pattern = r"初步评审|形式评审|资格评审"
        for match in re.finditer(prelim_pattern, text):
            start = max(0, match.start() - 50)
            end = min(len(text), match.end() + 200)
            context = text[start:end].replace("\n", " ").strip()
            method["preliminary_review"].append(context)

        return method

    def _extract_compliance_items(self, text: str) -> list:
        """提取合规项"""
        items = []

        # 资质要求
        qualification_pattern = r"资质[：:]\s*([^\n]+)"
        for match in re.finditer(qualification_pattern, text):
            items.append({
                "type": "QUALIFICATION",
                "requirement": match.group(1).strip(),
                "status": "PENDING"
            })

        return items

    def _extract_disqualification_items(self, text: str) -> list:
        """提取废标项"""
        items = []

        # 常见的废标条件
        patterns = [
            r"投标文件.+(?:逾期|超时)",
            r"(?:项目经理|技术负责人).+不在场",
            r"投标保证金.+(?:未交|不足)",
            r"资质.+(?:不满足|不符合)",
        ]

        for pattern in patterns:
            for match in re.finditer(pattern, text):
                items.append({
                    "item": match.group(0),
                    "location": f"位置{match.start()}",
                    "severity": "HIGH"
                })

        return items