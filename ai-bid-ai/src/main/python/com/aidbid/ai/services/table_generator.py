"""表格自动生成服务"""
import io
import logging
from typing import Dict, Any, List, Optional
from datetime import datetime, timedelta
import json

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import Rectangle
import numpy as np

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class TableGenerator:
    """技术标表格生成器"""

    def __init__(self):
        self.default_font = "宋体"
        self.default_size = 10.5
        self.header_bg_color = "4472C4"  # 蓝色表头
        self.alternate_row_color = "D9E2F3"  # 交替行颜色

    def _set_cell_style(self, cell, text: str, bold: bool = False,
                         font_size: int = 10.5, bg_color: Optional[str] = None,
                         align: str = "CENTER", font_color: Optional[str] = None):
        """设置单元格样式"""
        cell.text = text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "CENTER" else (
            WD_ALIGN_PARAGRAPH.LEFT if align == "LEFT" else WD_ALIGN_PARAGRAPH.RIGHT
        )
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = self.default_font

        # 设置字体颜色
        if font_color:
            run.font.color.rgb = RGBColor.from_string(font_color)

        # 设置单元格背景色
        if bg_color:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg_color)
            cell._tc.get_or_add_tcPr().append(shading)

    def _apply_table_border(self, table):
        """应用表格边框样式"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '4472C4')
            tblBorders.append(border)

        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def generate_qualification_table(self, requirements: List[Dict[str, Any]]) -> List[List[str]]:
        """生成资质要求对照表

        Args:
            requirements: 资质要求列表，每项包含:
                - name: 资质名称
                - required: 是否必需
                - level: 资质等级
                - deadline: 有效期限
                - status: 我方状态 (已满足/部分满足/不满足/待核实)
                - remark: 备注

        Returns:
            表格数据二维数组
        """
        headers = ["资质名称", "要求等级", "必需/可选", "有效期限", "我方状态", "备注"]
        rows = []

        for req in requirements:
            status = req.get("status", "")
            status_display = status
            if "已满足" in status:
                status_display = f"✓ {status}"
            elif "不满足" in status:
                status_display = f"✗ {status}"
            elif "部分" in status:
                status_display = f"△ {status}"

            rows.append([
                req.get("name", ""),
                req.get("level", ""),
                req.get("required", "必需"),
                req.get("deadline", ""),
                status_display,
                req.get("remark", "")
            ])

        return [headers] + rows

    def generate_personnel_table(self, project_scope: Dict[str, Any]) -> List[List[str]]:
        """生成项目人员配置表

        Args:
            project_scope: 项目范围，包含:
                - positions: 岗位列表
                - requirements: 人员要求

        Returns:
            表格数据二维数组
        """
        headers = ["序号", "岗位名称", "人数", "专业要求", "资质要求", "主要职责", "备注"]
        rows = []

        positions = project_scope.get("positions", [])

        for idx, pos in enumerate(positions, 1):
            rows.append([
                str(idx),
                pos.get("name", ""),
                str(pos.get("count", 1)),
                pos.get("major", ""),
                pos.get("qualification", ""),
                pos.get("responsibility", ""),
                pos.get("remark", "")
            ])

        # 添加合计行
        if rows:
            total_count = sum(int(row[2]) for row in rows if row[2].isdigit())
            rows.append([
                "", "合计", str(total_count), "", "", "", ""
            ])

        return [headers] + rows

    def generate_schedule_table(self, timeline: Dict[str, Any]) -> List[List[str]]:
        """生成施工进度计划表

        Args:
            timeline: 时间计划，包含:
                - phases: 阶段列表

        Returns:
            表格数据二维数组
        """
        headers = ["序号", "工作阶段", "开始时间", "结束时间", "持续天数", "关键节点", "备注"]
        rows = []

        phases = timeline.get("phases", [])

        for idx, phase in enumerate(phases, 1):
            rows.append([
                str(idx),
                phase.get("name", ""),
                phase.get("start_date", ""),
                phase.get("end_date", ""),
                str(phase.get("duration", 0)),
                phase.get("milestone", ""),
                phase.get("remark", "")
            ])

        return [headers] + rows

    def generate_equipment_table(self, equipment_list: List[Dict[str, Any]]) -> List[List[str]]:
        """生成设备清单表

        Args:
            equipment_list: 设备列表，每项包含:
                - name: 设备名称
                - model: 规格型号
                - count: 数量
                - owner_or_lease: 自有/租赁
                - condition: 现状
                - deployment_location: 部署位置

        Returns:
            表格数据二维数组
        """
        headers = ["序号", "设备名称", "规格型号", "数量", "来源", "现状", "部署位置", "备注"]
        rows = []

        for idx, eq in enumerate(equipment_list, 1):
            rows.append([
                str(idx),
                eq.get("name", ""),
                eq.get("model", ""),
                str(eq.get("count", 1)),
                eq.get("owner_or_lease", "自有"),
                eq.get("condition", "完好"),
                eq.get("deployment_location", ""),
                eq.get("remark", "")
            ])

        # 添加合计行
        if rows:
            total_count = sum(int(row[3]) for row in rows if row[3].isdigit())
            rows.append([
                "", "合计", "", str(total_count), "", "", "", ""
            ])

        return [headers] + rows

    def generate_safety_table(self, measures: List[Dict[str, Any]]) -> List[List[str]]:
        """生成安全措施检查表

        Args:
            measures: 安全措施列表，每项包含:
                - category: 类别
                - measure: 措施名称
                - standard: 检查标准
                - person_responsible: 责任人
                - check_frequency: 检查频率

        Returns:
            表格数据二维数组
        """
        headers = ["序号", "类别", "安全措施", "检查标准", "责任人", "检查频率", "备注"]
        rows = []

        for idx, m in enumerate(measures, 1):
            rows.append([
                str(idx),
                m.get("category", ""),
                m.get("measure", ""),
                m.get("standard", ""),
                m.get("person_responsible", ""),
                m.get("check_frequency", ""),
                m.get("remark", "")
            ])

        return [headers] + rows

    def generate_gantt_chart(self, timeline: Dict[str, Any],
                             output_format: str = "png",
                             width: int = 12, height: int = 8) -> bytes:
        """生成施工进度甘特图

        Args:
            timeline: 时间计划，包含 phases
            output_format: 输出格式 (png/svg)
            width: 图表宽度（英寸）
            height: 图表高度（英寸）

        Returns:
            图表二进制数据
        """
        phases = timeline.get("phases", [])

        if not phases:
            phases = [
                {"name": "准备阶段", "start_date": "2024-01-01", "end_date": "2024-01-15"},
                {"name": "施工阶段", "start_date": "2024-01-16", "end_date": "2024-03-15"},
                {"name": "验收阶段", "start_date": "2024-03-16", "end_date": "2024-03-31"},
            ]

        # 解析日期并计算持续时间
        tasks = []
        all_dates = []

        for phase in phases:
            name = phase.get("name", "")
            start = phase.get("start_date", "")
            end = phase.get("end_date", "")

            if start and end:
                start_dt = datetime.strptime(start, "%Y-%m-%d") if isinstance(start, str) else start
                end_dt = datetime.strptime(end, "%Y-%m-%d") if isinstance(end, str) else end
                duration = (end_dt - start_dt).days + 1

                tasks.append({
                    "name": name,
                    "start": start_dt,
                    "end": end_dt,
                    "duration": duration
                })

                all_dates.append(start_dt)
                all_dates.append(end_dt)

        if not tasks:
            # 默认示例数据
            tasks = [
                {"name": "项目准备", "start": datetime(2024, 1, 1), "end": datetime(2024, 1, 15), "duration": 15},
                {"name": "基础施工", "start": datetime(2024, 1, 16), "end": datetime(2024, 2, 15), "duration": 31},
                {"name": "主体结构", "start": datetime(2024, 2, 16), "end": datetime(2024, 3, 15), "duration": 29},
                {"name": "设备安装", "start": datetime(2024, 3, 1), "end": datetime(2024, 3, 31), "duration": 31},
                {"name": "竣工验收", "start": datetime(2024, 4, 1), "end": datetime(2024, 4, 15), "duration": 15},
            ]
            all_dates = [t["start"] for t in tasks] + [t["end"] for t in tasks]

        # 计算日期范围
        min_date = min(all_dates)
        max_date = max(all_dates)
        date_range = (max_date - min_date).days + 1

        # 设置中文字体
        plt.rcParams['font.sans-serif'] = ['WenQuanYi Micro Hei', 'Noto Sans CJK SC', 'AR PL UMing CN', 'SimHei', 'DejaVu Sans', 'Arial Unicode MS', 'sans-serif']
        plt.rcParams['axes.unicode_minus'] = False
        plt.rcParams['font.family'] = 'sans-serif'

        # 创建图表
        fig, ax = plt.subplots(figsize=(width, height))

        # 颜色方案
        colors = ['#4472C4', '#70AD47', '#ED7D31', '#FFC000', '#5B9BD5', '#7030A0', '#C00000', '#00B050']

        y_labels = []
        bar_positions = []

        for idx, task in enumerate(tasks):
            y_labels.append(task["name"])
            bar_positions.append(idx)

            # 计算条形位置和宽度
            start_offset = (task["start"] - min_date).days
            duration = task["duration"]

            # 绘制条形
            bar = Rectangle(
                (start_offset, idx - 0.4),
                duration,
                0.8,
                facecolor=colors[idx % len(colors)],
                edgecolor='white',
                linewidth=1
            )
            ax.add_patch(bar)

            # 添加工期天数标签
            ax.text(
                start_offset + duration / 2,
                idx,
                f"{duration}天",
                ha='center',
                va='center',
                fontsize=9,
                color='white',
                fontweight='bold'
            )

        # 设置Y轴
        ax.set_yticks(bar_positions)
        ax.set_yticklabels(y_labels, fontsize=10)
        ax.set_ylim(-0.6, len(tasks) - 0.4)

        # 设置X轴日期刻度
        date_ticks = []
        date_labels = []
        current = min_date
        while current <= max_date:
            date_ticks.append((current - min_date).days)
            date_labels.append(current.strftime("%m/%d"))
            current += timedelta(days=max(1, date_range // 15))

        ax.set_xticks(date_ticks)
        ax.set_xticklabels(date_labels, rotation=45, ha='right', fontsize=9)
        ax.set_xlim(-1, date_range + 1)

        # 添加网格线
        ax.grid(True, axis='x', linestyle='--', alpha=0.5)
        ax.set_axisbelow(True)

        # 设置标题
        project_name = timeline.get("project_name", "施工进度计划")
        ax.set_title(f"{project_name} - 施工进度计划（甘特图）",
                    fontsize=14, fontweight='bold', pad=15)

        # 添加图例
        legend_patches = [
            mpatches.Patch(color=colors[i % len(colors)], label=tasks[i]["name"])
            for i in range(min(5, len(tasks)))
        ]
        ax.legend(handles=legend_patches, loc='upper right', fontsize=8)

        # 添加今日线（如果当前日期在范围内）
        today = datetime.now()
        if min_date <= today <= max_date:
            today_offset = (today - min_date).days
            ax.axvline(x=today_offset, color='red', linestyle='--', linewidth=1.5, label='今日')
            ax.text(today_offset, len(tasks) - 0.3, '今日', color='red', fontsize=8, ha='center')

        # 调整布局
        plt.tight_layout()

        # 保存到字节流
        buffer = io.BytesIO()
        if output_format == "svg":
            plt.savefig(buffer, format='svg', dpi=150, bbox_inches='tight')
        else:
            plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        buffer.seek(0)
        plt.close(fig)

        return buffer.read()

    def add_styled_table_to_doc(self, doc: Document, table_data: List[List[str]],
                                  title: Optional[str] = None,
                                  header_style: str = "BLUE"):
        """向Word文档添加带样式的表格

        Args:
            doc: Word文档对象
            table_data: 表格数据，第一行为表头
            title: 表格标题
            header_style: 表头样式 (BLUE/GREEN/GRAY)
        """
        if not table_data or len(table_data) < 2:
            return

        # 添加表格标题
        if title:
            para = doc.add_paragraph()
            run = para.add_run(title)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = self.default_font

        # 创建表格
        headers = table_data[0]
        data_rows = table_data[1:]

        rows = len(data_rows) + 1
        cols = len(headers)

        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        # 设置表头样式
        header_colors = {
            "BLUE": "4472C4",
            "GREEN": "70AD47",
            "GRAY": "595959"
        }
        header_color = header_colors.get(header_style, "4472C4")

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            self._set_cell_style(
                cell, header,
                bold=True,
                font_size=10.5,
                bg_color=header_color,
                font_color="FFFFFF"
            )

        # 设置数据行样式
        for row_idx, row_data in enumerate(data_rows):
            row = table.rows[row_idx + 1]
            is_total_row = row_data[0] == "" and "合计" in row_data[1]

            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                bg = None

                # 合计行特殊样式
                if is_total_row:
                    bg = "D9E2F3"

                self._set_cell_style(
                    cell, cell_text,
                    bold=is_total_row,
                    font_size=10,
                    bg_color=bg,
                    align="CENTER" if col_idx > 1 else "LEFT"
                )

        # 应用边框
        self._apply_table_border(table)

        return table

    def export_table_to_word(self, table_data: List[List[str]],
                              title: str = "表格",
                              header_style: str = "BLUE") -> bytes:
        """导出表格为Word文档

        Args:
            table_data: 表格数据
            title: 表格标题
            header_style: 表头样式

        Returns:
            Word文档二进制数据
        """
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.name = self.default_font
        style.font.size = Pt(self.default_size)

        # 添加文档标题
        doc.add_heading(title, level=1)

        # 添加表格
        self.add_styled_table_to_doc(doc, table_data, header_style=header_style)

        # 保存到字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def export_gantt_to_word(self, gantt_image: bytes,
                              title: str = "施工进度计划",
                              doc_title: str = "施工进度计划（甘特图）") -> bytes:
        """将甘特图导出为Word文档

        Args:
            gantt_image: 甘特图二进制数据
            title: 甘特图标题
            doc_title: 文档标题

        Returns:
            Word文档二进制数据
        """
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.name = self.default_font
        style.font.size = Pt(self.default_size)

        # 添加文档标题
        doc.add_heading(doc_title, level=1)

        # 添加甘特图
        buffer = io.BytesIO(gantt_image)
        doc.add_picture(buffer, width=Inches(6.5))

        # 添加说明
        para = doc.add_paragraph()
        run = para.add_run(f"注：{title}")
        run.font.size = Pt(9)
        run.italic = True

        # 保存
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()


def create_sample_data():
    """创建示例数据用于测试"""
    return {
        "qualification_requirements": [
            {"name": "建筑业企业资质证书", "level": "二级及以上", "required": "必需",
             "deadline": "有效期内", "status": "已满足", "remark": "证书编号: XXX"},
            {"name": "安全生产许可证", "level": "有效期至2025年", "required": "必需",
             "deadline": "有效期内", "status": "已满足", "remark": ""},
            {"name": "ISO9001质量管理体系", "level": "认证通过", "required": "可选",
             "deadline": "有效期内", "status": "已满足", "remark": ""},
            {"name": "项目经理资质", "level": "一级建造师", "required": "必需",
             "deadline": "有效期内", "status": "部分满足", "remark": "现有人员符合要求"},
        ],
        "personnel_config": {
            "positions": [
                {"name": "项目经理", "count": 1, "major": "建筑工程",
                 "qualification": "一级建造师", "responsibility": "项目整体管理", "remark": ""},
                {"name": "技术负责人", "count": 1, "major": "土木工程",
                 "qualification": "中级职称", "responsibility": "技术管理", "remark": ""},
                {"name": "安全员", "count": 2, "major": "安全工程",
                 "qualification": "安全员证书", "responsibility": "安全管理", "remark": ""},
                {"name": "施工员", "count": 3, "major": "建筑工程",
                 "qualification": "施工员证书", "responsibility": "现场施工", "remark": ""},
                {"name": "质量员", "count": 1, "major": "质量管理",
                 "qualification": "质量员证书", "responsibility": "质量检查", "remark": ""},
            ]
        },
        "schedule_timeline": {
            "project_name": "某建设工程项目",
            "phases": [
                {"name": "项目准备", "start_date": "2024-01-01", "end_date": "2024-01-15",
                 "duration": 15, "milestone": "完成施工准备", "remark": ""},
                {"name": "基础施工", "start_date": "2024-01-16", "end_date": "2024-02-15",
                 "duration": 31, "milestone": "基础验收", "remark": ""},
                {"name": "主体结构", "start_date": "2024-02-16", "end_date": "2024-03-15",
                 "duration": 29, "milestone": "结构封顶", "remark": ""},
                {"name": "设备安装", "start_date": "2024-03-01", "end_date": "2024-03-31",
                 "duration": 31, "milestone": "安装调试", "remark": "与主体施工交叉进行"},
                {"name": "竣工验收", "start_date": "2024-04-01", "end_date": "2024-04-15",
                 "duration": 15, "milestone": "竣工验收", "remark": ""},
            ]
        },
        "equipment_list": [
            {"name": "塔式起重机", "model": "QTZ63", "count": 2, "owner_or_lease": "自有",
             "condition": "完好", "deployment_location": "施工现场", "remark": ""},
            {"name": "混凝土搅拌机", "model": "JS1000", "count": 1, "owner_or_lease": "自有",
             "condition": "完好", "deployment_location": "搅拌站", "remark": ""},
            {"name": "施工升降机", "model": "SC200/200", "count": 2, "owner_or_lease": "租赁",
             "condition": "完好", "deployment_location": "施工现场", "remark": ""},
            {"name": "挖掘机", "model": "卡特320", "count": 1, "owner_or_lease": "租赁",
             "condition": "完好", "deployment_location": "基础施工区", "remark": ""},
            {"name": "钢筋加工设备", "model": "GW40", "count": 2, "owner_or_lease": "自有",
             "condition": "完好", "deployment_location": "钢筋加工区", "remark": ""},
        ],
        "safety_measures": [
            {"category": "用电安全", "measure": "三级配电两级保护",
             "standard": "符合JGJ46-2005标准", "person_responsible": "电工班长",
             "check_frequency": "每日", "remark": ""},
            {"category": "用电安全", "measure": "漏电保护器配置",
             "standard": "额定漏电动作电流≤30mA", "person_responsible": "电工班长",
             "check_frequency": "每周", "remark": ""},
            {"category": "高空作业", "measure": "安全带佩戴",
             "standard": "高处作业必须系安全带", "person_responsible": "安全员",
             "check_frequency": "每作业前", "remark": ""},
            {"category": "高空作业", "measure": "脚手架验收",
             "standard": "符合JGJ130标准", "person_responsible": "技术负责人",
             "check_frequency": "每层验收", "remark": ""},
            {"category": "消防安全", "measure": "消防器材配置",
             "standard": "按面积每50㎡配灭火器1个", "person_responsible": "安全员",
             "check_frequency": "每月检查", "remark": ""},
            {"category": "机械安全", "measure": "塔吊定期检查",
             "standard": "符合GB5144标准", "person_responsible": "机械工程师",
             "check_frequency": "每月", "remark": ""},
        ]
    }