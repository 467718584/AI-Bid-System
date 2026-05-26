"""图片服务 - 图文并茂功能核心"""
import io
import logging
import os
import base64
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from enum import Enum
import httpx

logger = logging.getLogger(__name__)

# ============================================================
# 数据模型
# ============================================================

class ChartType(str, Enum):
    """图表类型"""
    BAR = "bar"           # 柱状图
    LINE = "line"         # 折线图
    PIE = "pie"           # 饼图
    TABLE = "table"       # 表格


@dataclass
class ImageResult:
    """图片检索结果"""
    image_id: str
    image_path: str
    image_url: Optional[str] = None
    caption: Optional[str] = None
    relevance_score: float = 0.0
    tags: List[str] = None

    def __post_init__(self):
        if self.tags is None:
            self.tags = []


@dataclass
class ChartSpec:
    """图表规格"""
    chart_type: ChartType
    title: str
    data: List[List[Any]]
    headers: Optional[List[str]] = None
    width: int = 6       # 宽度（英寸）
    height: int = 4      # 高度（英寸）


# ============================================================
# 图片服务
# ============================================================

class ImageService:
    """图片服务 - 支持图片检索、插入和图表生成"""

    def __init__(self, material_service_url: str = "http://localhost:8083"):
        self.material_service_url = material_service_url
        self._http_client: Optional[httpx.AsyncClient] = None
        self._local_image_cache: Dict[str, str] = {}  # image_id -> local_path

    @property
    def http_client(self) -> httpx.AsyncClient:
        if self._http_client is None:
            self._http_client = httpx.AsyncClient(timeout=30.0)
        return self._http_client

    async def close(self):
        if self._http_client:
            await self._http_client.aclose()
            self._http_client = None

    async def search_relevant_images(
        self,
        query: str,
        count: int = 3,
        project_id: Optional[str] = None
    ) -> List[ImageResult]:
        """搜索相关图片

        Args:
            query: 查询文本（关键词、描述等）
            count: 返回数量
            project_id: 项目ID（可选，用于限定素材库范围）

        Returns:
            图片结果列表
        """
        try:
            # 优先从本地素材库检索
            url = f"{self.material_service_url}/api/material/image/search"
            params = {
                "query": query,
                "count": count,
            }
            if project_id:
                params["projectId"] = project_id

            response = await self.http_client.get(url, params=params)

            if response.status_code == 200:
                result = response.json()
                images = []
                for item in result.get("data", {}).get("images", []):
                    images.append(ImageResult(
                        image_id=item.get("id", ""),
                        image_path=item.get("path", item.get("url", "")),
                        image_url=item.get("url"),
                        caption=item.get("caption", item.get("description", "")),
                        relevance_score=item.get("score", 0.0),
                        tags=item.get("tags", [])
                    ))
                return images
            else:
                logger.warning(f"Material service returned {response.status_code}")

        except Exception as e:
            logger.warning(f"Failed to search from material service: {e}")

        # 降级：使用关键词模拟返回占位图片
        return self._generate_placeholder_images(query, count)

    def _generate_placeholder_images(
        self,
        query: str,
        count: int
    ) -> List[ImageResult]:
        """生成占位图片（当素材库无可用图片时）"""
        results = []
        # 使用公共图库占位URL（实际项目中替换为真实素材库地址）
        placeholder_urls = [
            f"https://via.placeholder.com/400x300/e8f4f8/1a73e8?text={query[:10]}",
            f"https://via.placeholder.com/400x300/f8f4e8/ea4335?text={query[:10]}",
            f"https://via.placeholder.com/400x300/f4e8f8/34a853?text={query[:10]}",
        ]

        for i in range(min(count, len(placeholder_urls))):
            results.append(ImageResult(
                image_id=f"placeholder_{i}",
                image_path="",
                image_url=placeholder_urls[i],
                caption=f"相关素材：{query}",
                relevance_score=0.5,
                tags=[query]
            ))
        return results

    async def download_image(self, url: str) -> Optional[bytes]:
        """下载图片"""
        try:
            response = await self.http_client.get(url)
            response.raise_for_status()
            return response.content
        except Exception as e:
            logger.error(f"Failed to download image from {url}: {e}")
            return None

    async def insert_image_to_document(
        self,
        doc,
        image_source: Any,  # ImageResult | bytes | str(path/url)
        position: Optional[str] = None,
        width: Optional[float] = 6.0,
        height: Optional[float] = 4.0,
        caption: Optional[str] = None
    ) -> bool:
        """插入图片到Word文档

        Args:
            doc: python-docx Document对象
            image_source: 图片源（ImageResult对象、bytes数据或URL/path字符串）
            position: 插入位置描述（用于日志）
            width: 图片宽度（英寸）
            height: 图片高度（英寸）
            caption: 图片说明文字

        Returns:
            是否插入成功
        """
        try:
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # 获取图片数据
            image_bytes = await self._resolve_image_source(image_source)
            if image_bytes is None:
                logger.warning(f"Cannot resolve image source: {image_source}")
                return False

            # 转为 BytesIO
            image_stream = io.BytesIO(image_bytes)

            # 插入图片
            run = doc.add_paragraph().add_run()
            run.add_picture(image_stream, width=Inches(width), height=Inches(height))

            # 添加图片说明
            if caption:
                caption_para = doc.add_paragraph()
                caption_para.add_run(caption).italic = True
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            logger.info(f"Image inserted at {position or 'end of document'}")
            return True

        except Exception as e:
            logger.error(f"Failed to insert image: {e}")
            return False

    async def _resolve_image_source(self, source: Any) -> Optional[bytes]:
        """解析图片源为字节数据"""
        if isinstance(source, ImageResult):
            if source.image_path and os.path.exists(source.image_path):
                with open(source.image_path, "rb") as f:
                    return f.read()
            elif source.image_url:
                return await self.download_image(source.image_url)
            return None
        elif isinstance(source, bytes):
            return source
        elif isinstance(source, str):
            if source.startswith("http://") or source.startswith("https://"):
                return await self.download_image(source)
            elif os.path.exists(source):
                with open(source, "rb") as f:
                    return f.read()
            return None
        return None

    # ============================================================
    # 图表生成
    # ============================================================

    async def generate_chart_from_data(
        self,
        data: List[List[Any]],
        chart_type: str = "bar",
        title: str = "",
        headers: Optional[List[str]] = None
    ) -> Optional[bytes]:
        """从数据生成图表图片

        Args:
            data: 二维数据列表
            chart_type: 图表类型 (bar/line/pie)
            title: 图表标题
            headers: 表头（第一行是否作为表头）

        Returns:
            图表PNG图片的字节数据
        """
        try:
            import matplotlib
            matplotlib.use("Agg")  # 无头模式
            import matplotlib.pyplot as plt
            import numpy as np

            if not data or not data[0]:
                return None

            # 提取数据
            headers_list = headers if headers else (data[0] if len(data) > 1 else None)
            data_rows = data[1:] if headers_list else data

            if chart_type == "bar":
                return self._generate_bar_chart(data_rows, headers_list, title)
            elif chart_type == "line":
                return self._generate_line_chart(data_rows, headers_list, title)
            elif chart_type == "pie":
                return self._generate_pie_chart(data_rows, headers_list, title)
            else:
                # 降级返回表格图片
                return self._generate_table_image(data, title)

        except ImportError as e:
            logger.warning(f"matplotlib not available: {e}")
            return None
        except Exception as e:
            logger.error(f"Chart generation failed: {e}")
            return None

    def _generate_bar_chart(
        self,
        data: List[List[Any]],
        labels: Optional[List[str]],
        title: str
    ) -> Optional[bytes]:
        """生成柱状图"""
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            import numpy as np

            if not labels:
                labels = [f"Item {i+1}" for i in range(len(data))]

            values = []
            for row in data:
                try:
                    values.append(float(row[1] if len(row) > 1 else row[0]))
                except (ValueError, TypeError):
                    values.append(float(row[0]))

            plt.figure(figsize=(6, 4))
            bars = plt.bar(range(len(values)), values, color=["#4285f4", "#34a853", "#fbbc05", "#ea4335"])
            plt.xlabel(labels[1] if len(labels) > 1 else "")
            plt.ylabel(labels[2] if len(labels) > 2 else "")
            plt.title(title)
            plt.xticks(range(len(data)), labels[:len(values)], rotation=45)
            plt.tight_layout()

            buf = io.BytesIO()
            plt.savefig(buf, format="png", dpi=150)
            plt.close()
            buf.seek(0)
            return buf.read()
        except Exception as e:
            logger.error(f"Bar chart failed: {e}")
            return None

    def _generate_line_chart(
        self,
        data: List[List[Any]],
        labels: Optional[List[str]],
        title: str
    ) -> Optional[bytes]:
        """生成折线图"""
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            import numpy as np

            if not labels:
                labels = [f"Item {i+1}" for i in range(len(data))]

            values = [float(row[1] if len(row) > 1 else row[0]) for row in data]

            plt.figure(figsize=(6, 4))
            plt.plot(range(len(values)), values, marker="o", color="#4285f4", linewidth=2)
            plt.xlabel(labels[1] if len(labels) > 1 else "")
            plt.ylabel(labels[2] if len(labels) > 2 else "")
            plt.title(title)
            plt.xticks(range(len(data)), labels[:len(values)], rotation=45)
            plt.grid(True, linestyle="--", alpha=0.6)
            plt.tight_layout()

            buf = io.BytesIO()
            plt.savefig(buf, format="png", dpi=150)
            plt.close()
            buf.seek(0)
            return buf.read()
        except Exception as e:
            logger.error(f"Line chart failed: {e}")
            return None

    def _generate_pie_chart(
        self,
        data: List[List[Any]],
        labels: Optional[List[str]],
        title: str
    ) -> Optional[bytes]:
        """生成饼图"""
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt

            if not labels:
                labels = [f"Item {i+1}" for i in range(len(data))]

            values = [float(row[1] if len(row) > 1 else row[0]) for row in data]

            plt.figure(figsize=(5, 5))
            colors = ["#4285f4", "#34a853", "#fbbc05", "#ea4335", "#9c27b0", "#00bcd4"]
            plt.pie(values, labels=labels, colors=colors[:len(values)],
                    autopct="%1.1f%%", startangle=90)
            plt.title(title)
            plt.tight_layout()

            buf = io.BytesIO()
            plt.savefig(buf, format="png", dpi=150)
            plt.close()
            buf.seek(0)
            return buf.read()
        except Exception as e:
            logger.error(f"Pie chart failed: {e}")
            return None

    def _generate_table_image(
        self,
        data: List[List[Any]],
        title: str
    ) -> Optional[bytes]:
        """生成表格图片"""
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            import numpy as np

            if not data:
                return None

            fig, ax = plt.subplots(figsize=(8, max(1, len(data) * 0.5)))
            ax.axis("tight")

            table = ax.table(
                cellText=data,
                colLabels=data[0] if len(data) > 0 else None,
                cellLoc="center",
                loc="center"
            )
            table.auto_set_font_size(False)
            table.set_fontsize(10)
            table.scale(1.2, 1.5)

            # 表头样式
            for (row, col), cell in table.get_celld().items():
                if row == 0:
                    cell.set_facecolor("#4285f4")
                    cell.set_text_props(color="white", weight="bold")

            ax.axis("off")
            if title:
                ax.set_title(title, fontsize=12, weight="bold")
            plt.tight_layout()

            buf = io.BytesIO()
            plt.savefig(buf, format="png", dpi=150)
            plt.close()
            buf.seek(0)
            return buf.read()
        except Exception as e:
            logger.error(f"Table image failed: {e}")
            return None

    async def generate_styled_table(
        self,
        doc,
        data: List[List[Any]],
        headers: Optional[List[str]] = None,
        title: Optional[str] = None,
        styled: bool = True
    ):
        """向Word文档插入带样式的表格

        Args:
            doc: python-docx Document对象
            data: 二维数据
            headers: 表头列表
            title: 表格标题
            styled: 是否美化样式
        """
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # 表格标题
        if title:
            caption_para = doc.add_paragraph()
            run = caption_para.add_run(title)
            run.bold = True
            run.font.size = Pt(11)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        rows = len(data) + (1 if headers else 0)
        cols = len(data[0]) if data else 0

        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        # 表头
        if headers:
            header_row = table.rows[0]
            for i, h in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = h
                para = cell.paragraphs[0]
                run = para.runs[0]
                run.bold = True
                run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if styled:
                    self._set_cell_shading(cell, "4472C4")

        # 数据行
        start = 1 if headers else 0
        for i, row_data in enumerate(data):
            row = table.rows[i + start]
            for j, text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(text)
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                if styled:
                    self._set_cell_shading(cell, "FFFFFF" if i % 2 == 0 else "F2F2F2")

        return table

    def _set_cell_shading(self, cell, fill_color: str):
        """设置单元格背景色"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), fill_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)


# ============================================================
# 全局单例
# ============================================================

_image_service: Optional[ImageService] = None


def get_image_service() -> ImageService:
    """获取图片服务实例"""
    global _image_service
    if _image_service is None:
        material_url = os.getenv("MATERIAL_SERVICE_URL", "http://localhost:8083")
        _image_service = ImageService(material_service_url=material_url)
    return _image_service


async def close_image_service():
    """关闭图片服务"""
    global _image_service
    if _image_service:
        await _image_service.close()
        _image_service = None