"""Word文档导出模块"""
import io
import re
import json
import logging
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class DocumentExporter:
    """Word文档导出器"""

    def __init__(self):
        self.default_font = "宋体"
        self.default_size = 12

    def create_document(self) -> Document:
        """创建新文档"""
        doc = Document()
        # 设置默认样式
        style = doc.styles["Normal"]
        style.font.name = self.default_font
        style.font.size = Pt(self.default_size)
        return doc

    def add_title(self, doc: Document, title: str, level: int = 1):
        """添加标题"""
        heading = doc.add_heading(title, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_paragraph(self, doc: Document, text: str, bold: bool = False,
                      font_size: Optional[int] = None, alignment: str = "LEFT"):
        """添加段落"""
        para = doc.add_paragraph()
        para.add_run(text).bold = bold

        if font_size:
            para.runs[0].font.size = Pt(font_size)

        if alignment == "CENTER":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "RIGHT":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        return para

    def add_table(self, doc: Document, data: List[List[str]],
                  headers: Optional[List[str]] = None):
        """添加表格"""
        if not data:
            return

        rows = len(data) + (1 if headers else 0)
        cols = len(data[0]) if data else 0

        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        # 添加表头
        if headers:
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加数据
        start_idx = 1 if headers else 0
        for i, row_data in enumerate(data):
            row = table.rows[i + start_idx]
            for j, cell_text in enumerate(row_data):
                row.cells[j].text = str(cell_text)

    def export_analysis_report(self, analysis_result: Dict[str, Any]) -> bytes:
        """导出分析报告为Word文档"""
        doc = self.create_document()

        # 标题
        self.add_title(doc, "招标文件智能分析报告", level=0)

        # 基本信息
        if "basic_info" in analysis_result:
            self.add_title(doc, "一、基本信息", level=1)
            basic_info = analysis_result["basic_info"]
            for key, value in basic_info.items():
                self.add_paragraph(doc, f"{key}: {value}")

        # 评分方法
        if "scoring_method" in analysis_result:
            self.add_title(doc, "二、评分方法", level=1)
            scoring = analysis_result["scoring_method"]
            if "preliminary_review" in scoring:
                self.add_paragraph(doc, "初步评审:", bold=True)
                for item in scoring["preliminary_review"]:
                    self.add_paragraph(doc, f"  - {item}")

        # 合规检查项
        if "compliance_items" in analysis_result:
            self.add_title(doc, "三、合规检查项", level=1)
            for item in analysis_result["compliance_items"]:
                status = item.get("status", "UNKNOWN")
                self.add_paragraph(doc, f"[{status}] {item.get('requirement', '')}")

        # 废标条款
        if "disqualification_items" in analysis_result:
            self.add_title(doc, "四、废标条款", level=1)
            for item in analysis_result["disqualification_items"]:
                self.add_paragraph(doc, f"- {item.get('item', '')}")

        # AI分析结论
        if "analysis" in analysis_result:
            self.add_title(doc, "五、AI分析结论", level=1)
            self.add_paragraph(doc, analysis_result["analysis"])

        # 保存到字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def export_compliance_report(self, compliance_data: Dict[str, Any]) -> bytes:
        """导出合规检查报告"""
        doc = self.create_document()

        self.add_title(doc, "投标合规性检查报告", level=0)

        if "project_name" in compliance_data:
            self.add_paragraph(doc, f"项目名称: {compliance_data['project_name']}")

        if "check_items" in compliance_data:
            self.add_title(doc, "检查项目明细", level=1)

            table_data = []
            for item in compliance_data["check_items"]:
                table_data.append([
                    item.get("name", ""),
                    item.get("status", ""),
                    item.get("suggestion", "")
                ])

            if table_data:
                self.add_table(doc, table_data,
                               headers=["检查项", "状态", "建议"])

        # 保存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def export_bid_summary(self, summary_data: Dict[str, Any]) -> bytes:
        """导出投标摘要"""
        doc = self.create_document()

        self.add_title(doc, "投标文件摘要", level=0)

        # 项目信息
        self.add_title(doc, "项目信息", level=1)
        project_info = summary_data.get("project_info", {})
        for key, value in project_info.items():
            self.add_paragraph(doc, f"{key}: {value}")

        # 资质匹配
        self.add_title(doc, "资质匹配情况", level=1)
        qualifications = summary_data.get("qualifications", [])
        for q in qualifications:
            status = "通过" if q.get("matched") else "不通过"
            self.add_paragraph(doc, f"[{status}] {q.get('name', '')}")

        # 保存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _set_cell_shading(self, cell, fill_color: str):
        """设置单元格背景色"""
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), fill_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # ============================================================
    # 图片插入方法
    # ============================================================

    def add_image(
        self,
        doc: Document,
        image_data: bytes,
        caption: Optional[str] = None,
        width: float = 6.0,
        height: Optional[float] = None
    ):
        """插入图片到文档

        Args:
            doc: Document对象
            image_data: 图片字节数据
            caption: 图片说明
            width: 宽度（英寸）
            height: 高度（英寸），None则自动按比例
        """
        image_stream = io.BytesIO(image_data)
        para = doc.add_paragraph()
        run = para.add_run()

        if height:
            run.add_picture(image_stream, width=Inches(width), height=Inches(height))
        else:
            run.add_picture(image_stream, width=Inches(width))

        if caption:
            caption_para = doc.add_paragraph()
            caption_run = caption_para.add_run(caption)
            caption_run.italic = True
            caption_run.font.size = Pt(9)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_image_from_url(
        self,
        doc: Document,
        image_url: str,
        caption: Optional[str] = None,
        width: float = 6.0
    ):
        """从URL插入图片"""
        import httpx
        try:
            response = httpx.get(image_url, timeout=15)
            response.raise_for_status()
            self.add_image(doc, response.content, caption, width)
        except Exception as e:
            logger.warning(f"Failed to add image from URL {image_url}: {e}")
            # 降级：添加占位文本
            para = doc.add_paragraph()
            run = para.add_run(f"[图片: {image_url}]")
            run.italic = True

    # ============================================================
    # Markdown解析导出
    # ============================================================

    def add_markdown_content(
        self,
        doc: Document,
        markdown_text: str,
        include_images: bool = True,
        include_tables: bool = True
    ):
        """将Markdown内容解析并添加到Word文档

        支持以下元素：
        - 标题 (# / ## / ###)
        - 段落
        - 列表 (- / 1. )
        - 表格 (| ... |)
        - 图表占位符 (**[图表: ...]**)
        - 链接和图片引用
        """
        import re

        lines = markdown_text.split("\n")
        in_table = False
        table_lines = []

        for line in lines:
            stripped = line.strip()

            # 跳过空行
            if not stripped:
                if in_table and table_lines:
                    # 解析并添加表格
                    self._add_table_from_lines(doc, table_lines)
                    table_lines = []
                    in_table = False
                continue

            # 检测表格行
            if stripped.startswith("|"):
                in_table = True
                table_lines.append(stripped)
                continue
            else:
                # 结束表格
                if in_table and table_lines:
                    self._add_table_from_lines(doc, table_lines)
                    table_lines = []
                    in_table = False

            # 标题
            if stripped.startswith("### "):
                self.add_title(doc, stripped[4:], level=3)
            elif stripped.startswith("## "):
                self.add_title(doc, stripped[3:], level=2)
            elif stripped.startswith("# "):
                self.add_title(doc, stripped[2:], level=1)

            # 图表占位符
            elif re.match(r"^\*\*\[图表:\s*.+\]\*\*$", stripped):
                chart_title = stripped.strip("**[]").replace("图表: ", "")
                para = doc.add_paragraph()
                run = para.add_run(f"[图表: {chart_title}]")
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 无序列表
            elif stripped.startswith("- ") or stripped.startswith("* "):
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(stripped[2:])

            # 有序列表
            elif re.match(r"^\d+\.\s+", stripped):
                para = doc.add_paragraph(style="List Number")
                para.add_run(re.sub(r"^\d+\.\s+", "", stripped))

            # 普通段落
            else:
                # 清理Markdown格式后添加
                clean_text = self._clean_markdown_text(stripped)
                if clean_text:
                    self.add_paragraph(doc, clean_text)

        # 处理末尾表格
        if in_table and table_lines:
            self._add_table_from_lines(doc, table_lines)

    def _add_table_from_lines(
        self,
        doc: Document,
        table_lines: List[str]
    ):
        """从Markdown表格行创建Word表格"""
        if not table_lines:
            return

        # 解析表格数据
        rows_data = []
        headers = []

        for i, line in enumerate(table_lines):
            # 去除首尾|
            cells = [c.strip() for c in line.strip("|").split("|")]

            # 分隔行跳过
            if all(c in ("---", ":---", ":-:", "---:") for c in cells):
                headers = rows_data[0] if rows_data else []
                rows_data = rows_data[1:] if rows_data else []
                continue

            if i == 0:
                headers = cells
            else:
                rows_data.append(cells)

        if not headers and not rows_data:
            return

        self.add_styled_table(doc, rows_data, headers if headers else None)

    def add_styled_table(
        self,
        doc: Document,
        data: List[List[Any]],
        headers: Optional[List[str]] = None,
        title: Optional[str] = None
    ):
        """添加带样式的表格"""
        if not data:
            return

        # 表格标题
        if title:
            para = doc.add_paragraph()
            run = para.add_run(title)
            run.bold = True
            run.font.size = Pt(11)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        rows_count = len(data) + (1 if headers else 0)
        cols_count = len(data[0]) if data else 0

        table = doc.add_table(rows=rows_count, cols=cols_count)
        table.style = "Table Grid"

        # 表头
        if headers:
            header_row = table.rows[0]
            for i, h in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = str(h)
                para = cell.paragraphs[0]
                run = para.runs[0]
                run.bold = True
                run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._set_cell_shading(cell, "4472C4")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # 数据行
        start = 1 if headers else 0
        for i, row_data in enumerate(data):
            row = table.rows[i + start]
            for j, text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(text)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                # 斑马条纹
                fill = "FFFFFF" if i % 2 == 0 else "F2F2F2"
                self._set_cell_shading(cell, fill)

    def _clean_markdown_text(self, text: str) -> str:
        """清理Markdown格式"""
        import re
        # 粗体 **text**
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        # 斜体 *text*
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        # 行内代码 `code`
        text = re.sub(r"`(.+?)`", r"\1", text)
        # 链接 [text](url)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        # 图片 ![alt](url)
        text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
        # 保留表格图片标记（已在上方处理）
        text = re.sub(r"!\[\s\S]*?\]\(table:[\s\S]*?\)", "", text)
        text = re.sub(r"!\[\s\S]*?\]\(chart:[\s\S]*?\)", "", text)
        return text.strip()