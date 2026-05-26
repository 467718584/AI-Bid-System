#!/usr/bin/env python3
"""
MD to DOCX 转换脚本
使用模板将markdown文件批量转换为docx
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 配置
WORKSPACE = "/home/zzy/.openclaw/workspace/workspace-bid"
TEMPLATE_PATH = os.path.join(WORKSPACE, "template.docx")
OUTPUT_DIR = os.path.join(WORKSPACE, "docs-docx")
DOCS_DIR = os.path.join(WORKSPACE, "docs")

# 样式映射
STYLE_MAP = {
    'h1': 'Heading 1',
    'h2': 'Heading 2',
    'h3': 'Heading 3',
    'h4': 'Heading 4',
    'h5': 'Heading 5',
    'h6': 'Heading 6',
}

def read_markdown(file_path):
    """读取markdown文件"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def extract_frontmatter(content):
    """提取YAML前置matter"""
    match = re.match(r'^---\n(.*?)\n---\n', content, re.DOTALL)
    if match:
        return content[len(match.group(0)):]
    return content

def md_to_docx(content, doc):
    """将markdown内容转换为docx"""
    
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        # 标题处理 (# ## ### 等)
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            style_name = f'Heading {level}'
            p = doc.add_paragraph(text, style=style_name)
            i += 1
            continue
        
        # 表格处理
        if '|' in line and line.strip().startswith('|'):
            # 收集表格行
            table_lines = [line]
            while i + 1 < len(lines) and '|' in (lines[i+1] if lines[i+1].strip() else ''):
                i += 1
                table_lines.append(lines[i])
            
            # 解析表格
            if len(table_lines) >= 2:
                # 过滤分隔行
                data_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:|]+\|$', l)]
                if data_lines:
                    table = doc.add_table(rows=len(data_lines), cols=len(data_lines[0].split('|')))
                    table.style = 'Table Grid'
                    for ri, row_data in enumerate(data_lines):
                        cells = row_data.split('|')[1:-1]
                        for ci, cell in enumerate(cells):
                            table.rows[ri].cells[ci].text = cell.strip()
            i += 1
            continue
        
        # 列表处理 (- 或 * 或数字.)
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            marker = list_match.group(2)
            text = list_match.group(3).strip()
            
            # 处理嵌套列表
            p = doc.add_paragraph(style='List Bullet' if marker in ['-', '*'] else 'List Number')
            p.paragraph_format.left_indent = Cm(indent * 0.5 + 0.5)
            p.add_run(text)
            i += 1
            continue
        
        # 代码块处理
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            # 添加代码块样式
            p = doc.add_paragraph('\n'.join(code_lines), style='Compact')
            i += 1
            continue
        
        # 分隔线
        if re.match(r'^---+$', line.strip()):
            i += 1
            continue
        
        # 强调处理 (**bold** 或 *italic*)
        # 先简单处理为普通文本
        p = doc.add_paragraph(line, style='Body Text')
        i += 1

def convert_file(md_path, output_dir, template_path):
    """转换单个文件"""
    # 读取markdown
    content = read_markdown(md_path)
    content = extract_frontmatter(content)
    
    # 打开模板
    doc = Document(template_path)
    
    # 清空模板中的示例段落（保留样式）
    # 获取所有paragraphs，除了前面几个示例标题
    paras = doc.paragraphs
    # 保留前10个段落（样式示例），删除其余的
    for p in paras[10:]:
        p._element.getparent().remove(p._element)
    
    # 转换内容
    md_to_docx(content, doc)
    
    # 保存
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    output_name = Path(md_path).stem + '.docx'
    output_path = os.path.join(output_dir, output_name)
    doc.save(output_path)
    
    return output_path

def main():
    """主函数"""
    # 创建输出目录
    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
    
    # 获取所有md文件
    md_files = []
    for root, dirs, files in os.walk(DOCS_DIR):
        for f in files:
            if f.endswith('.md'):
                md_files.append(os.path.join(root, f))
    
    # 按文件名排序
    md_files.sort()
    
    print(f"找到 {len(md_files)} 个md文件")
    print(f"模板: {TEMPLATE_PATH}")
    print(f"输出目录: {OUTPUT_DIR}")
    print()
    
    success = []
    failed = []
    
    for md_file in md_files:
        try:
            rel_path = os.path.relpath(md_file, DOCS_DIR)
            output_path = convert_file(md_file, OUTPUT_DIR, TEMPLATE_PATH)
            success.append(rel_path)
            print(f"✅ {rel_path}")
        except Exception as e:
            failed.append((rel_path, str(e)))
            print(f"❌ {rel_path}: {e}")
    
    print()
    print(f"转换完成: {len(success)} 成功, {len(failed)} 失败")
    
    if failed:
        print("\n失败文件:")
        for path, error in failed:
            print(f"  - {path}: {error}")

if __name__ == '__main__':
    main()